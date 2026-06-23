"""Resume parsing (spec Section 3.1).

Extracts skills, tools/languages, coursework, project keywords, degree/major, and
prior experience titles from a PDF/DOCX/TXT resume, then builds a field-agnostic
weighted keyword map used by the matcher.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from . import domain
from .models import ResumeProfile

log = logging.getLogger("internfinder.resume")


# ----------------------------------------------------------------- text extract
def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_docx(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume type: {suffix} (use PDF, DOCX, or TXT)")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise ImportError("pdfplumber is required for PDF resumes: pip install pdfplumber") from exc

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    if not text:
        log.warning(
            "No text extracted from %s — it may be a scanned/image PDF. "
            "Skills extraction will be empty; export a text-based PDF or DOCX.",
            path.name,
        )
    return text


def _extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ImportError("python-docx is required for DOCX resumes: pip install python-docx") from exc

    document = docx.Document(str(path))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:  # skills are often laid out in tables
        for row in table.rows:
            lines.extend(cell.text for cell in row.cells)
    return "\n".join(lines).strip()


# --------------------------------------------------------------- section split
_SECTION_HEADERS = {
    "skills": re.compile(r"^\s*(technical skills|skills|technical proficiencies|technologies)\b", re.I),
    "education": re.compile(r"^\s*(education|academics)\b", re.I),
    "experience": re.compile(r"^\s*(experience|work experience|employment|professional experience)\b", re.I),
    "projects": re.compile(r"^\s*(projects|technical projects|personal projects|academic projects)\b", re.I),
    "coursework": re.compile(r"^\s*(coursework|relevant coursework|courses)\b", re.I),
}


def _split_sections(text: str) -> dict[str, str]:
    """Bucket lines into known sections by header detection. Best-effort."""
    sections: dict[str, list[str]] = {k: [] for k in _SECTION_HEADERS}
    sections["_preamble"] = []
    current = "_preamble"
    for line in text.splitlines():
        matched = None
        for name, pat in _SECTION_HEADERS.items():
            if pat.match(line):
                matched = name
                break
        if matched:
            current = matched
            # keep any trailing content on the same line (e.g. "Skills: C, Python")
            remainder = line.split(":", 1)[1] if ":" in line else ""
            if remainder.strip():
                sections[current].append(remainder)
            continue
        sections[current].append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items()}


# ------------------------------------------------------------------- extractors
_DEGREE_RE = re.compile(
    r"\b(b\.?s\.?|b\.?eng\.?|bachelor(?:'s)?|m\.?s\.?|master(?:'s)?|ph\.?d\.?|associate)\b",
    re.I,
)
_MAJOR_RE = re.compile(
    r"\b(?:in|of)\s+([A-Za-z][A-Za-z&/ ]{3,60}?"
    r"(?:engineering|science|physics|mathematics|math|business|finance|economics|"
    r"marketing|design|biology|chemistry|nursing|healthcare|communications?|"
    r"psychology|policy|political science|public health|accounting|education|"
    r"statistics|data science|computer science))\b",
    re.I,
)
_TITLE_RE = re.compile(
    r"\b("
    r"(?:senior|junior|lead|staff|principal|associate|assistant|research|product|project|program|"
    r"marketing|brand|growth|sales|business|finance|financial|accounting|investment|operations|"
    r"strategy|design|ux|ui|graphic|content|social|communications|public relations|policy|legal|"
    r"data|analytics|science|biology|chemistry|clinical|lab|nursing|healthcare|education|"
    r"software|hardware|firmware|embedded|electrical|mechanical|systems|test|validation)\s+"
    r"(?:intern(?:ship)?|analyst|associate|assistant|coordinator|manager|designer|researcher|"
    r"scientist|technician|engineer(?:ing)?|developer|specialist|consultant)"
    r"|(?:marketing|finance|business|product|design|research|operations|sales|policy|clinical|"
    r"engineering|software|hardware|data)\s+intern"
    r"|intern(?:ship)?|co-?op|teaching assistant|research assistant"
    r")\b",
    re.I,
)


def _extract_degree_major(edu_text: str, full_text: str) -> tuple[str, str]:
    scope = edu_text or full_text
    degree = ""
    dm = _DEGREE_RE.search(scope)
    if dm:
        degree = dm.group(1).upper().replace(".", "")
    major = ""
    mm = _MAJOR_RE.search(scope)
    if mm:
        major = re.sub(r"\s+", " ", mm.group(1)).strip()
    return degree, major


def _extract_titles(exp_text: str, full_text: str) -> list[str]:
    titles: list[str] = []
    seen: set[str] = set()
    for m in _TITLE_RE.finditer(exp_text or full_text):
        t = re.sub(r"\s+", " ", m.group(0)).strip().title()
        key = t.lower()
        if key not in seen:
            seen.add(key)
            titles.append(t)
    return titles[:12]


def _extract_name(preamble: str) -> str:
    for line in preamble.splitlines():
        line = line.strip()
        # First short line of 2-4 capitalized words, no digits/@, is usually the name.
        if line and "@" not in line and not any(ch.isdigit() for ch in line):
            words = line.split()
            if 1 < len(words) <= 4 and all(w[:1].isupper() for w in words if w):
                return line
    return ""


# --------------------------------------------------------------------- public
def parse_resume(path: str | Path, config: dict | None = None) -> ResumeProfile:
    config = config or {}
    text = extract_text(path)
    sections = _split_sections(text)

    skills_text = sections.get("skills", "")
    projects_text = sections.get("projects", "")
    experience_text = sections.get("experience", "")
    # Tech terms anywhere in the resume, but skills/projects sections are richest.
    all_terms = domain.extract_known_terms(text)
    skill_terms = domain.extract_known_terms(skills_text + "\n" + projects_text)

    tools_languages = [t for t in all_terms if t in domain.TOOLS_AND_LANGUAGES]
    degree, major = _extract_degree_major(sections.get("education", ""), text)
    titles = _extract_titles(experience_text, text)
    coursework = domain.extract_known_terms(sections.get("coursework", ""))
    projects = domain.extract_known_terms(projects_text)

    # Build weighted keyword map: synonym-expanded from everything we found, then
    # boosted for terms that appeared in the dedicated skills section, then the
    # user's configured priority keywords pinned high.
    weights = domain.expand_terms(set(all_terms))
    for t in skill_terms:
        weights[t] = max(weights.get(t, 0.0), 1.4)

    # --- FIELD-AGNOSTIC layer (Section 3.1, generalized) ------------------
    # Pull plain-language skills from the resume so non-engineering candidates
    # (finance, biology, design, marketing, nursing, …) match on *their* terms,
    # not just a hardware lexicon. Skills/projects/experience are the strongest
    # signal; the rest of the resume contributes weakly.
    for term, w in domain.extract_generic_terms(
        "\n".join([skills_text, projects_text, experience_text])
    ).items():
        weights[term] = max(weights.get(term, 0.0), 0.6 + 0.6 * w)  # ≈0.9–1.2
    for term, w in domain.extract_generic_terms(text).items():
        weights[term] = max(weights.get(term, 0.0), w)              # ≈0.45–0.95

    priority = list(config.get("domain", {}).get("priority_keywords", []) if config else [])
    # The candidate's stated target role/field is the single strongest signal of
    # intent — pin it (and any synonyms) highest, above everything inferred.
    target_role = (config.get("search", {}).get("target_role", "") if config else "") or ""
    role_terms = [t.strip().lower() for t in re.split(r"[,/;]+", target_role) if t.strip()]
    for k in role_terms:
        weights[k] = max(weights.get(k, 0.0), 1.8)
        for syn in domain.SYNONYMS.get(k, ()):
            weights[syn] = max(weights.get(syn, 0.0), 0.9)
    for kw in priority:
        k = kw.strip().lower()
        if not k:
            continue
        weights[k] = max(weights.get(k, 0.0), 1.5)
        for syn in domain.SYNONYMS.get(k, ()):
            weights[syn] = max(weights.get(syn, 0.0), 0.8)

    profile = ResumeProfile(
        raw_text=text,
        name=_extract_name(sections.get("_preamble", "")),
        degree=degree,
        major=major,
        target_role=target_role.strip(),
        skills=sorted({t for t in all_terms}),
        tools_languages=sorted(set(tools_languages)),
        coursework=coursework,
        projects=projects,
        experience_titles=titles,
        weighted_keywords=dict(sorted(weights.items(), key=lambda kv: -kv[1])),
    )
    profile.summary = _build_summary(profile)
    log.info(
        "Parsed resume: %d skills, %d tools/langs, degree=%r major=%r, %d titles",
        len(profile.skills), len(profile.tools_languages), profile.degree,
        profile.major, len(profile.experience_titles),
    )
    return profile


def _build_summary(p: ResumeProfile) -> str:
    """Compact natural-language summary fed to the LLM matcher (Section 7)."""
    bits: list[str] = []
    who = " ".join(x for x in [p.degree, p.major] if x).strip()
    if who:
        bits.append(f"Candidate pursuing {who}.")
    if p.target_role:
        bits.append(f"Target roles: {p.target_role}.")
    if p.tools_languages:
        bits.append("Tools/languages: " + ", ".join(p.tools_languages[:18]) + ".")
    other_skills = [s for s in p.skills if s not in p.tools_languages]
    if other_skills:
        bits.append("Skills/concepts: " + ", ".join(other_skills[:25]) + ".")
    if p.coursework:
        bits.append("Coursework: " + ", ".join(p.coursework[:12]) + ".")
    if p.experience_titles:
        bits.append("Prior roles: " + ", ".join(p.experience_titles[:8]) + ".")
    return " ".join(bits)
